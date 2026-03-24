"""
Script para gerar o modelo Word editável (.docx).
Executa uma vez: python gerar_exemplo_word.py
O ficheiro gerado vai para frontend/assets/documento_exemplo.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = os.path.join(
    os.path.dirname(__file__),
    "..", "frontend", "assets", "documento_exemplo.docx"
)

AZUL      = RGBColor(0x00, 0x33, 0x66)
AZUL_CLARO = RGBColor(0xE6, 0xED, 0xFF)
BRANCO    = RGBColor(0xFF, 0xFF, 0xFF)
CINZA_LINHA = RGBColor(0xF4, 0xF8, 0xFF)


def cor_celula(celula, rgb: RGBColor):
    """Define cor de fundo de uma célula."""
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def borda_celula(celula, tamanho=4, cor="CCCCCC"):
    """Define bordas finas numa célula."""
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(tamanho))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), cor)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def adicionar_topico(doc, numero, titulo, linhas):
    """Adiciona uma secção numerada com tabela de 2 colunas."""

    # Cabeçalho da secção
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"{numero}. {titulo}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL

    # Tabela
    tabela = doc.add_table(rows=1, cols=2)
    tabela.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabela.style = "Table Grid"

    # Larguras das colunas (aprox. 8,5 cm + 9,5 cm = 18 cm)
    for i, largura in enumerate([Cm(8.5), Cm(9.5)]):
        for row in tabela.rows:
            row.cells[i].width = largura

    # Cabeçalho da tabela
    cabecalhos = ["Parâmetro", "Valor"]
    for i, texto in enumerate(cabecalhos):
        celula = tabela.rows[0].cells[i]
        cor_celula(celula, AZUL)
        borda_celula(celula, tamanho=4, cor="003366")
        p_cel = celula.paragraphs[0]
        p_cel.paragraph_format.space_before = Pt(2)
        p_cel.paragraph_format.space_after  = Pt(2)
        run_cel = p_cel.add_run(f"  {texto}")
        run_cel.bold = True
        run_cel.font.color.rgb = BRANCO
        run_cel.font.size = Pt(9)

    # Linhas de dados
    for idx, (label, valor) in enumerate(linhas):
        row = tabela.add_row()
        cor_fundo = CINZA_LINHA if idx % 2 == 1 else BRANCO

        # Coluna esquerda (label a negrito)
        celula_esq = row.cells[0]
        celula_esq.width = Cm(8.5)
        cor_celula(celula_esq, cor_fundo)
        borda_celula(celula_esq)
        p_esq = celula_esq.paragraphs[0]
        p_esq.paragraph_format.space_before = Pt(1)
        p_esq.paragraph_format.space_after  = Pt(1)
        run_esq = p_esq.add_run(f"  {label}")
        run_esq.bold = True
        run_esq.font.size = Pt(9)

        # Coluna direita (valor normal)
        celula_dir = row.cells[1]
        celula_dir.width = Cm(9.5)
        cor_celula(celula_dir, cor_fundo)
        borda_celula(celula_dir)
        p_dir = celula_dir.paragraphs[0]
        p_dir.paragraph_format.space_before = Pt(1)
        p_dir.paragraph_format.space_after  = Pt(1)
        run_dir = p_dir.add_run(f"  {valor}")
        run_dir.font.size = Pt(9)

    doc.add_paragraph()  # espaço entre secções


def gerar():
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Título principal
    titulo = doc.add_paragraph()
    titulo.paragraph_format.space_after = Pt(4)
    run_titulo = titulo.add_run("Ficha Técnica de Equipamento de Impressão")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    run_titulo.font.color.rgb = AZUL

    # Linha separadora sob o título
    p_linha = doc.add_paragraph()
    p_linha.paragraph_format.space_before = Pt(0)
    p_linha.paragraph_format.space_after  = Pt(6)
    pPr = p_linha._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "003366")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Tabela de cabeçalho (info do equipamento)
    t_info = doc.add_table(rows=2, cols=2)
    t_info.alignment = WD_TABLE_ALIGNMENT.LEFT
    t_info.style = "Table Grid"
    dados_info = [
        ("Marca", "Konica Minolta",                      "Modelo", "AccurioPress C4080"),
        ("Tipo",  "Impressora Digital de Produção a Cores", "Data / Ref.", "24/03/2026 — KM-APC4080-2026"),
    ]
    for r_idx, (k1, v1, k2, v2) in enumerate(dados_info):
        row = t_info.rows[r_idx]
        for c_idx, (chave, val) in enumerate([(k1, v1), (k2, v2)]):
            celula = row.cells[c_idx]
            cor_celula(celula, AZUL_CLARO)
            borda_celula(celula, tamanho=4, cor="003366")
            p_cel = celula.paragraphs[0]
            p_cel.paragraph_format.space_before = Pt(2)
            p_cel.paragraph_format.space_after  = Pt(2)
            run_k = p_cel.add_run(f"{chave}: ")
            run_k.bold = True
            run_k.font.size = Pt(9)
            run_k.font.color.rgb = AZUL
            run_v = p_cel.add_run(val)
            run_v.font.size = Pt(9)
    doc.add_paragraph()

    # ── Secções ────────────────────────────────────────────────────────────────
    adicionar_topico(doc, 1, "Velocidade de Impressão", [
        ("Impressão a cores (A4)",            "80 ppm (páginas por minuto)"),
        ("Impressão a preto e branco (A4)",   "80 ppm"),
        ("Impressão duplex automático (A4)",  "60 ppm"),
        ("Formato SRA3",                      "40 ppm"),
    ])
    adicionar_topico(doc, 2, "Qualidade de Impressão", [
        ("Resolução máxima",    "3600 × 2400 dpi"),
        ("Tecnologia",          "Electrofotografia a laser"),
        ("Reprodução de cor",   "CMYK + Toner especial opcional"),
        ("Profundidade de cor", "8 bits por canal"),
    ])
    adicionar_topico(doc, 3, "Capacidade de Papel", [
        ("Capacidade total de entrada", "8.300 folhas"),
        ("Gramagem suportada",          "52 g/m² até 400 g/m²"),
        ("Formatos suportados",         "A6 a SRA3 (320 × 450 mm)"),
        ("Papel couchê",                "Sim"),
        ("Envelopes e etiquetas",       "Sim"),
    ])
    adicionar_topico(doc, 4, "Características Gerais, Acabamentos e Desempenho", [
        ("Impressão frente e verso (duplex)", "Sim, automático"),
        ("Conectividade",                     "Ethernet 1 Gbps, USB 3.0"),
        ("Protocolo de impressão",            "PCL6, PostScript 3, PDF 1.7"),
        ("Sistemas operativos compatíveis",   "Windows, macOS, Linux"),
        ("Agrafagem automática",              "Sim (até 100 folhas)"),
        ("Furação",                           "Sim (2 e 4 furos)"),
        ("Dobragem",                          "Sim (em Z, cavalete, simples)"),
        ("Volume mensal recomendado",         "Até 850.000 páginas/mês"),
        ("Volume máximo suportado",           "1.500.000 páginas/mês"),
        ("Tempo de aquecimento",              "Menos de 25 segundos"),
        ("Tempo para primeira página",        "Menos de 5 segundos"),
        ("Consumo em funcionamento",          "2,5 kW"),
        ("Consumo em modo de espera",         "0,5 kW"),
        ("Nível de ruído (em funcionamento)", "65 dB(A)"),
        ("Certificações ambientais",          "Energy Star, EPEAT Gold"),
    ])

    # Rodapé
    p_rodape = doc.add_paragraph(
        "Documento de exemplo — edita os valores, adiciona ou remove tópicos e linhas conforme necessário. "
        "Guarda como PDF para utilizar no validador."
    )
    p_rodape.paragraph_format.space_before = Pt(8)
    run_rodape = p_rodape.runs[0]
    run_rodape.font.size = Pt(8)
    run_rodape.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_rodape.italic = True

    doc.save(OUTPUT)
    print(f"Word gerado: {os.path.abspath(OUTPUT)}")


if __name__ == "__main__":
    gerar()
